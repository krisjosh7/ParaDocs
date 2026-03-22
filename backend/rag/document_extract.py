from __future__ import annotations

import logging
from io import BytesIO
from pathlib import Path

try:
    from PIL import Image
except ImportError:
    Image = None  # type: ignore[misc, assignment]

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None  # type: ignore[misc, assignment]

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore[misc, assignment]

logger = logging.getLogger(__name__)

_IMAGE_MIME: dict[str, str] = {
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".gif": "image/gif",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
}

# Groq base64 payload limit ~4MB; stay under ~3.5MB raw before encoding.
_MAX_VISION_IMAGE_BYTES = 2_800_000


def is_image_suffix(suffix: str) -> bool:
    return suffix.lower() in _IMAGE_MIME


# Extensions we treat as audio/video-for-transcription for context RAG (ElevenLabs STT).
_AUDIO_SUFFIXES: frozenset[str] = frozenset(
    {
        ".mp3",
        ".wav",
        ".m4a",
        ".aac",
        ".flac",
        ".ogg",
        ".opus",
        ".webm",
        ".wma",
        ".mp2",
        ".amr",
        ".3gp",
        ".aiff",
        ".aif",
        ".mp4",  # often contains spoken audio; ElevenLabs accepts video
        ".mpeg",
        ".mpg",
    },
)


def is_audio_suffix(suffix: str) -> bool:
    return suffix.lower() in _AUDIO_SUFFIXES


def prepare_image_for_vision_api(path: Path) -> tuple[bytes, str] | None:
    """
    Return (bytes, mime_type) for a vision API data URL.
    Re-encodes as JPEG when needed so payload stays under provider limits.
    """
    try:
        raw = path.read_bytes()
    except OSError as e:
        logger.warning("Could not read image %s: %s", path, e)
        return None

    suffix = path.suffix.lower()

    if Image is not None:
        try:
            img = Image.open(BytesIO(raw))
            img.load()
            if getattr(img, "n_frames", 1) > 1:
                img.seek(0)
            rgb = img.convert("RGB")
            quality = 88
            scale = 1.0
            w, h = rgb.size
            for _ in range(24):
                if scale < 1.0:
                    nw = max(1, int(w * scale))
                    nh = max(1, int(h * scale))
                    resized = rgb.resize((nw, nh), Image.Resampling.LANCZOS)
                else:
                    resized = rgb
                buf = BytesIO()
                resized.save(buf, format="JPEG", quality=quality, optimize=True)
                jpeg_bytes = buf.getvalue()
                if len(jpeg_bytes) <= _MAX_VISION_IMAGE_BYTES:
                    return jpeg_bytes, "image/jpeg"
                if quality > 55:
                    quality -= 7
                else:
                    scale *= 0.82
            logger.warning("Could not compress image enough for vision API: %s", path)
            return None
        except Exception as e:
            logger.warning("PIL could not prepare image %s: %s", path, e)

    if len(raw) <= _MAX_VISION_IMAGE_BYTES:
        mime = _IMAGE_MIME.get(suffix)
        if mime:
            return raw, mime
    logger.warning("Image too large for vision API (install Pillow for auto-resize): %s", path)
    return None


def extract_text_from_pdf(path: Path) -> str:
    """Extract plain text from a PDF file. Returns empty string on failure."""
    if PdfReader is None:
        logger.warning("pypdf not installed; cannot extract PDF text")
        return ""

    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            try:
                t = page.extract_text()
            except Exception:
                t = None
            if t and t.strip():
                parts.append(t.strip())
        return "\n\n".join(parts).strip()
    except Exception as e:
        logger.warning("PDF text extraction failed for %s: %s", path, e)
        return ""


def extract_text_from_docx(path: Path) -> str:
    """Extract plain text from a .docx file. Returns empty string on failure."""
    if DocxDocument is None:
        logger.warning("python-docx not installed; cannot extract DOCX text")
        return ""

    try:
        doc = DocxDocument(str(path))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts).strip()
    except Exception as e:
        logger.warning("DOCX text extraction failed for %s: %s", path, e)
        return ""
