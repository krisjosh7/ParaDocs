from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from rag.document_extract import extract_text_from_docx, extract_text_from_pdf


def test_extract_docx_paragraphs(tmp_path: Path) -> None:
    p = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("First line")
    doc.add_paragraph("Second line")
    doc.save(p)

    text = extract_text_from_docx(p)
    assert "First line" in text
    assert "Second line" in text


def test_extract_pdf_mocked(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    p = tmp_path / "x.pdf"
    p.write_bytes(b"%PDF-1.4 fake")

    class FakePage:
        def extract_text(self) -> str:
            return "Hello from PDF"

    class FakeReader:
        def __init__(self, path: str) -> None:
            self.pages = [FakePage()]

    monkeypatch.setattr("rag.document_extract.PdfReader", FakeReader)

    assert "Hello from PDF" in extract_text_from_pdf(p)


def test_context_rag_pdf_includes_caption_and_extracted(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("CASES_ROOT", str(tmp_path))
    case_id = "c1"
    base = tmp_path / case_id / "contexts" / "files"
    base.mkdir(parents=True)
    pdf_path = base / "ctx-1.pdf"
    pdf_path.write_bytes(b"x")

    monkeypatch.setattr("rag.context_rag.extract_text_from_pdf", lambda p: "Body from PDF")

    from rag.context_rag import build_raw_text_for_context_rag

    row = {
        "id": "ctx-1",
        "type": "document",
        "title": "Exhibit",
        "caption": "Key exhibit for damages",
        "stored_file": "ctx-1.pdf",
    }
    out = build_raw_text_for_context_rag(case_id, row)
    assert "Key exhibit for damages" in out
    assert "Body from PDF" in out
    assert "Extracted PDF text" in out


def test_context_rag_image_includes_caption_and_vision_text(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("CASES_ROOT", str(tmp_path))
    case_id = "c1"
    base = tmp_path / case_id / "contexts" / "files"
    base.mkdir(parents=True)
    png_path = base / "ctx-2.png"
    png_path.write_bytes(b"\x89PNG\r\n\x1a\n")  # minimal header; vision is mocked

    def fake_describe(path: Path, *, caption: str = "") -> str:
        assert path == png_path
        return f"Model description (caption was {caption!r})"

    monkeypatch.setattr("rag.context_rag.describe_image_for_ingest", fake_describe)

    from rag.context_rag import build_raw_text_for_context_rag

    row = {
        "id": "ctx-2",
        "type": "document",
        "title": "Photo exhibit",
        "caption": "Accident scene",
        "stored_file": "ctx-2.png",
    }
    out = build_raw_text_for_context_rag(case_id, row)
    assert "Accident scene" in out
    assert "Model description" in out
    assert "Visual description (from image analysis)" in out


def test_context_rag_audio_includes_caption_and_transcript(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("CASES_ROOT", str(tmp_path))
    case_id = "c1"
    base = tmp_path / case_id / "contexts" / "files"
    base.mkdir(parents=True)
    audio_path = base / "ctx-3.mp3"
    audio_path.write_bytes(b"ID3fake")

    def fake_transcribe(path: Path) -> str:
        assert path == audio_path
        return "This is the spoken content from the recording."

    monkeypatch.setattr("rag.context_rag.transcribe_audio_for_ingest", fake_transcribe)

    from rag.context_rag import build_raw_text_for_context_rag

    row = {
        "id": "ctx-3",
        "type": "audio",
        "title": "Deposition clip",
        "caption": "Witness statement excerpt",
        "stored_file": "ctx-3.mp3",
    }
    out = build_raw_text_for_context_rag(case_id, row)
    assert "Witness statement excerpt" in out
    assert "spoken content" in out
    assert "Transcript (ElevenLabs speech-to-text)" in out
